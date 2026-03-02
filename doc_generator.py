# doc_generator.py
# =================
# This module's ONLY job: take PR data + AI analysis and produce a .docx file.
#
# Uses python-docx (pure Python, no Node.js needed).
# Install with: pip install python-docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx2pdf import convert
from docx.oxml import OxmlElement
from collections import Counter
import os
from datetime import datetime


# ── Colour palette (RGB tuples) ──────────────────────────
PRIMARY    = RGBColor(0x1E, 0x3A, 0x5F)   # dark blue
DARK       = RGBColor(0x1F, 0x29, 0x37)   # near black
GRAY       = RGBColor(0x6B, 0x72, 0x80)   # medium gray
LIGHT_GRAY = RGBColor(0x9C, 0xA3, 0xAF)   # light gray
GREEN      = RGBColor(0x05, 0x96, 0x69)   # green
PURPLE     = RGBColor(0x7C, 0x3A, 0xED)   # purple
AMBER      = RGBColor(0xD9, 0x74, 0x06)   # amber
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xEF, 0xF6, 0xFF)   # pale blue background
LIGHT_GREEN= RGBColor(0xEC, 0xFD, 0xF5)   # pale green background


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC"):
    """Add a light border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_paragraph(doc, text, size=11, bold=False, italic=False,
                   color=None, align=None, space_before=0, space_after=6):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_left_bar_paragraph(doc, text, bar_color="3B82F6", bg_color="EFF6FF"):
    """Add a paragraph with a coloured left border — used for AI analysis and summaries."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.3)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), bar_color)
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def _add_divider(doc):
    """Add a thin horizontal line between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# BUILDERS — one function per document section
# ─────────────────────────────────────────────

def _build_title(doc, repo, generated_at):
    """Title block at the top."""
    _add_paragraph(doc, "Pull Request Report",
                   size=26, bold=True, color=PRIMARY,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=10, space_after=4)
    _add_paragraph(doc, f"Repository: {repo}",
                   size=13, color=GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_paragraph(doc, f"Generated: {generated_at}",
                   size=10, italic=True, color=LIGHT_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    _add_divider(doc)


def _build_overview(doc, prs):
    """Stats row: total PRs, merged, contributors, lines changed."""
    _add_paragraph(doc, "Overview", size=16, bold=True,
                   color=PRIMARY, space_before=12, space_after=8)

    merged       = sum(1 for pr in prs if pr.get("merged_at") != "Not merged")
    contributors = len(set(pr["author"] for pr in prs))
    additions    = sum(pr.get("additions", 0) for pr in prs)
    deletions    = sum(pr.get("deletions", 0) for pr in prs)

    stats = [
        (str(len(prs)),               "Total PRs",     "1E3A5F", "EEF2FF"),
        (str(merged),                 "Merged",         "059669", "ECFDF5"),
        (str(contributors),           "Contributors",   "7C3AED", "F5F3FF"),
        (f"+{additions}/-{deletions}", "Lines Changed", "D97706", "FFFBEB"),
    ]

    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for col_idx, (value, label, text_color, bg_color) in enumerate(stats):
        val_cell = table.cell(0, col_idx)
        _set_cell_bg(val_cell, bg_color)
        _set_cell_border(val_cell)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_run = val_cell.paragraphs[0].add_run(value)
        val_run.bold = True
        val_run.font.size = Pt(18)
        r, g, b = int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16)
        val_run.font.color.rgb = RGBColor(r, g, b)

        lbl_cell = table.cell(1, col_idx)
        _set_cell_bg(lbl_cell, bg_color)
        _set_cell_border(lbl_cell)
        lbl_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_run = lbl_cell.paragraphs[0].add_run(label)
        lbl_run.font.size = Pt(9)
        lbl_run.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _build_label_breakdown(doc, prs):
    """
    Label breakdown section — shows count of each label across all PRs.
    Helps manager see at a glance how many bugs, enhancements, etc. were worked on.
    Only shown if at least one PR has a label.
    """
    # Flatten all labels from all PRs into one list
    all_labels = [label for pr in prs for label in pr.get("labels", [])]

    if not all_labels:
        return  # Skip section entirely if no labels exist

    label_counts = Counter(all_labels)  # e.g. {"bug": 3, "enhancement": 2}

    _add_paragraph(doc, "Labels Breakdown", size=16, bold=True,
                   color=PRIMARY, space_before=12, space_after=8)

    # Colour mapping for known GitHub labels
    label_colors = {
        "bug":              ("FEE2E2", "DC2626"),   # red
        "enhancement":      ("DCFCE7", "059669"),   # green
        "documentation":    ("DBEAFE", "2563EB"),   # blue
        "duplicate":        ("FEF9C3", "D97706"),   # yellow
        "help wanted":      ("FFE4E6", "E11D48"),   # pink
        "good first issue": ("F3E8FF", "7C3AED"),   # purple
        "invalid":          ("F3F4F6", "6B7280"),   # gray
        "question":         ("E0F2FE", "0284C7"),   # light blue
        "wontfix":          ("F9FAFB", "374151"),   # light gray
    }

    table = doc.add_table(rows=2, cols=len(label_counts))
    table.style = "Table Grid"

    for col_idx, (label, count) in enumerate(label_counts.items()):
        bg, text_color = label_colors.get(label.lower(), ("F3F4F6", "374151"))
        r, g, b = int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16)

        # Row 0 — count number
        count_cell = table.cell(0, col_idx)
        _set_cell_bg(count_cell, bg)
        _set_cell_border(count_cell)
        count_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        count_run = count_cell.paragraphs[0].add_run(str(count))
        count_run.bold = True
        count_run.font.size = Pt(18)
        count_run.font.color.rgb = RGBColor(r, g, b)

        # Row 1 — label name
        label_cell = table.cell(1, col_idx)
        _set_cell_bg(label_cell, bg)
        _set_cell_border(label_cell)
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    _add_divider(doc)


def _build_ai_analysis(doc, overall_analysis):
    """AI analysis blockquote section."""
    _add_paragraph(doc, "AI Analysis", size=16, bold=True,
                   color=PRIMARY, space_before=10, space_after=6)
    _add_left_bar_paragraph(doc, overall_analysis)
    _add_divider(doc)


def _build_pr_card(doc, pr):
    """
    Detail card for a single PR.
    Includes: info table, AI summary, description,
              linked issues, and code diffs.
    """
    # ── PR heading ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(f"PR #{pr['number']} — ")
    run1.bold = True
    run1.font.size = Pt(13)
    run1.font.color.rgb = GRAY
    run2 = p.add_run(pr["title"])
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = DARK

    # ── AI summary (green left bar) ──
    if pr.get("ai_summary"):
        _add_left_bar_paragraph(doc, f"💡 {pr['ai_summary']}",
                                bar_color="10B981", bg_color="ECFDF5")

    # ── Info table (2 columns: label | value) ──
    rows = [
        ("Author",            pr.get("author", "N/A")),
        ("Labels",            ", ".join(pr.get("labels") or []) or "None"),
        ("Branch",            f"{pr.get('head_branch', 'N/A')}  →  {pr.get('base_branch', 'N/A')}"),
        ("Status",            f"Merged on {pr['merged_at']}" if pr.get("merged_at") != "Not merged" else pr.get("state", "").upper()),
        ("Merged By",         pr.get("merged_by") or "N/A"),
        ("Created",           pr.get("created_at", "N/A")),
        ("Code Changes",      f"+{pr.get('additions', 0)} additions, -{pr.get('deletions', 0)} deletions"),
        ("Commits",           str(pr.get("commits", 0))),
        ("Approvals",         ", ".join(pr.get("approvals") or []) or "None"),
        ("Changes Requested", ", ".join(pr.get("changes_requested_by") or []) or "None"),
        ("Comments",          f"{pr.get('comments', 0)} general, {pr.get('review_comments', 0)} review"),
        ("Files Changed",     ", ".join((pr.get("files_changed") or [])[:5]) or "N/A"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)

        _set_cell_bg(label_cell, "F9FAFB")
        _set_cell_border(label_cell)
        _set_cell_border(value_cell)

        label_cell.width = Inches(1.8)
        value_cell.width = Inches(5.7)

        label_cell.paragraphs[0].clear()
        lr = label_cell.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = DARK

        value_cell.paragraphs[0].clear()
        vr = value_cell.paragraphs[0].add_run(str(value))
        vr.font.size = Pt(10)

    # ── PR description ──
    if pr.get("body") and pr["body"] != "No description":
        _add_paragraph(doc, "Description:", size=10, bold=True,
                       color=DARK, space_before=8, space_after=2)
        _add_paragraph(doc, pr["body"][:400], size=10,
                       color=GRAY, space_after=8)

    # ── Linked Issues ──
    # These are GitHub Issues referenced in the PR with "Fixes #42" or "Closes #123"
    # This shows what bug/feature request the PR was solving
    linked_issues = pr.get("linked_issues", [])
    _add_paragraph(doc, "Linked Issues:", size=10, bold=True,
                   color=DARK, space_before=8, space_after=4)

    if linked_issues:
        for issue in linked_issues:
            status = "Closed ✅" if issue["state"] == "closed" else "Still Open 🔴"
            _add_paragraph(doc,
                f"Issue #{issue['number']}: {issue['title']}",
                size=10, bold=True, color=PRIMARY, space_after=2)
            _add_paragraph(doc,
                f"  Reported by: {issue['author']}  |  Created: {issue['created_at']}  |  Status: {status}",
                size=9, color=GRAY, space_after=2)
            if issue.get("body") and issue["body"] != "No description":
                _add_paragraph(doc, f"  {issue['body'][:200]}",
                               size=9, italic=True, color=GRAY, space_after=6)
    else:
        _add_paragraph(doc,
            "No linked issues found — PR description did not reference any issues with Fixes/Closes #number",
            size=9, italic=True, color=LIGHT_GRAY, space_after=6)

    # ── Code Diffs ──
    # Shows exactly what lines were added (+) and removed (-) in each file
    file_diffs = pr.get("file_diffs", [])
    if file_diffs:
        _add_paragraph(doc, "Code Changes:", size=10, bold=True,
                       color=DARK, space_before=8, space_after=4)

        for diff in file_diffs:
            if not diff.get("patch"):
                continue  # skip binary files with no patch text

            status = diff.get("status", "modified")
            status_label = {
                "added":    "ADDED",
                "modified": "MODIFIED",
                "removed":  "REMOVED",
                "renamed":  "RENAMED"
            }.get(status, "CHANGED")

            # File name + status line
            _add_paragraph(doc,
                f"{status_label}: {diff.get('filename')}  "
                f"(+{diff.get('additions', 0)} / -{diff.get('deletions', 0)})",
                size=9, bold=True, color=PRIMARY,
                space_before=6, space_after=2)

            # The actual diff text — lines with + are additions, lines with - are removals
            _add_paragraph(doc,
                diff.get("patch", "")[:800],  # limit length to avoid huge pages
                size=8, color=GRAY, space_after=4)


def _build_footer(doc, generated_at):
    """Footer line at the bottom."""
    _add_divider(doc)
    _add_paragraph(doc,
        f"Report generated by PR Documentation Agent  •  {generated_at}",
        size=9, italic=True, color=LIGHT_GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)


# ─────────────────────────────────────────────
# PUBLIC FUNCTION — called by main.py
# ─────────────────────────────────────────────

def generate(pr_data_list, analysis, repo, output_path):
    """
    Generate a .docx report from PR data and AI analysis.

    Args:
        pr_data_list : list of clean PR dicts (from github_client.py)
        analysis     : AI analysis dict (from oci_client.py)
        repo         : "owner/repo" string (for the document title)
        output_path  : where to save the .docx file
    """
    print("\n📄 Generating Word document...")

    # Attach AI summaries to each PR
    summary_map = {
        str(s["number"]): s["summary"]
        for s in analysis.get("pr_summaries", [])
    }
    for pr in pr_data_list:
        pr["ai_summary"] = summary_map.get(str(pr["number"]), pr.get("body", "")[:200])

    generated_at = datetime.now().strftime("%B %d, %Y at %H:%M")

    doc = Document()

    # Set page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Build each section in order
    _build_title(doc, repo, generated_at)
    _build_overview(doc, pr_data_list)
    _build_label_breakdown(doc, pr_data_list)   # ← label breakdown after overview
    _build_ai_analysis(doc, analysis.get("overall_analysis", "No analysis available."))

    _add_paragraph(doc, "Pull Request Details", size=16, bold=True,
                   color=PRIMARY, space_before=10, space_after=6)

    for i, pr in enumerate(pr_data_list):
        _build_pr_card(doc, pr)
        if i < len(pr_data_list) - 1:
            _add_divider(doc)

    _build_footer(doc, generated_at)

    doc.save(output_path)
    print(f"✅ Word document saved: {output_path}")

    # Convert to PDF — saves in same folder as the .docx
    pdf_path = output_path.replace(".docx", ".pdf")
    try:
        convert(output_path, pdf_path)
        print(f"✅ PDF saved: {pdf_path}")
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")
        print("   Make sure Word is installed — docx2pdf needs Word to convert")