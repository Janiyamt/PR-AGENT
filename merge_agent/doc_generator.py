# merge_agent/doc_generator.py
# ==============================
# Generates a .docx report from all stored merge events.
# Called after every successful merge — rewrites the full document
# so it always reflects the complete history in merge-log.json.
#
# The document is structured as:
#   1. Cover / title block
#   2. Summary stats table (total merges, contributors, total lines changed)
#   3. One "merge card" per event, newest first
#   4. Footer
#
# Shares styling helpers with v1's doc_generator but is completely separate
# so both versions can coexist in the repo without conflicts.

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import Counter

# ── Colour palette ──────────────────────────────────────────────────────────
PRIMARY    = RGBColor(0x1E, 0x3A, 0x5F)
DARK       = RGBColor(0x1F, 0x29, 0x37)
GRAY       = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0x9C, 0xA3, 0xAF)
GREEN      = RGBColor(0x05, 0x96, 0x69)
PURPLE     = RGBColor(0x7C, 0x3A, 0xED)
AMBER      = RGBColor(0xD9, 0x74, 0x06)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

_LOG_DIR     = os.path.join(os.path.dirname(__file__), "logs")
_OUTPUT_FILE = os.path.join(_LOG_DIR, "merge-events-report.docx")


# ─────────────────────────────────────────────
# PUBLIC ENTRY POINT
# ─────────────────────────────────────────────

def generate(all_events: list, repo: str, output_path: str = _OUTPUT_FILE) -> None:
    """
    Generate (or overwrite) the merge events Word document.

    Args:
        all_events:  list of metadata dicts from log_store.load_all()
        repo:        "owner/repo" string used in the document header
        output_path: where to save the .docx
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    generated_at = datetime.now().strftime("%B %d, %Y at %H:%M UTC")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    _build_title(doc, repo, generated_at, len(all_events))
    _build_summary_stats(doc, all_events)
    _build_ai_overview(doc, all_events)

    _add_paragraph(doc, "Merge Event Details", size=16, bold=True,
                   color=PRIMARY, space_before=14, space_after=8)

    print(f"  Generating doc with {len(all_events)} events...")
    for e in all_events:
        print(f"     - PR #{e.get('pr_number')} {e.get('pr_title', '')}")

    for i, event in enumerate(all_events):
        try:
            _build_merge_card(doc, event)
        except Exception as e:
            print(f"  Skipped PR #{event.get('pr_number')} due to error: {e}")
        if i < len(all_events) - 1:
            _add_divider(doc)

    _build_footer(doc, generated_at)

    doc.save(output_path)
    print(f"  ✅ Merge events report saved: {output_path}")


# ─────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────

def _build_title(doc, repo, generated_at, event_count):
    _add_paragraph(doc, "Merge Events Report",
                   size=26, bold=True, color=PRIMARY,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=10, space_after=4)
    _add_paragraph(doc, f"Repository: {repo}",
                   size=13, color=GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_paragraph(doc, f"Generated: {generated_at}  •  {event_count} merge events on record",
                   size=10, italic=True, color=LIGHT_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    _add_divider(doc)


def _build_summary_stats(doc, events):
    """High-level stats across all recorded merge events."""
    if not events:
        return

    _add_paragraph(doc, "Summary", size=16, bold=True,
                   color=PRIMARY, space_before=12, space_after=8)

    total_additions  = sum(e.get("additions", 0) for e in events)
    total_deletions  = sum(e.get("deletions", 0) for e in events)
    contributors     = len(set(e.get("pr_author", "") for e in events))
    mergers          = len(set(e.get("merged_by", "") for e in events))

    stats = [
        (str(len(events)),                 "Total Merges",    "1E3A5F", "EEF2FF"),
        (str(contributors),                "PR Authors",      "059669", "ECFDF5"),
        (str(mergers),                     "Merged By",       "7C3AED", "F5F3FF"),
        (f"+{total_additions}/{total_deletions}", "Lines Δ", "D97706", "FFFBEB"),
    ]

    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for col_idx, (value, label, text_hex, bg_hex) in enumerate(stats):
        r, g, b = int(text_hex[0:2], 16), int(text_hex[2:4], 16), int(text_hex[4:6], 16)
        color   = RGBColor(r, g, b)

        val_cell = table.cell(0, col_idx)
        _set_cell_bg(val_cell, bg_hex)
        _set_cell_border(val_cell)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_run = val_cell.paragraphs[0].add_run(value)
        val_run.bold = True
        val_run.font.size = Pt(18)
        val_run.font.color.rgb = color

        lbl_cell = table.cell(1, col_idx)
        _set_cell_bg(lbl_cell, bg_hex)
        _set_cell_border(lbl_cell)
        lbl_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_run = lbl_cell.paragraphs[0].add_run(label)
        lbl_run.font.size = Pt(9)
        lbl_run.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    _add_divider(doc)


def _build_ai_overview(doc, events):
    """Collect and display AI summaries in a quick-scan section."""
    summaries = [
        (e.get("pr_number"), e.get("pr_title", ""), e.get("ai_analysis", {}).get("pr_summary", ""))
        for e in events
        if e.get("ai_analysis", {}).get("pr_summary")
    ]
    if not summaries:
        return

    _add_paragraph(doc, "AI Summaries", size=16, bold=True,
                   color=PRIMARY, space_before=12, space_after=6)

    for pr_num, pr_title, summary in summaries:
        _add_paragraph(doc, f"PR #{pr_num} — {pr_title}",
                       size=10, bold=True, color=DARK, space_after=2)
        _add_left_bar_paragraph(doc, summary)

    _add_divider(doc)


def _build_merge_card(doc, event):
    """Full detail card for a single merge event."""
    ai = event.get("ai_analysis", {})

    # Heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"PR #{event.get('pr_number')} — ")
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = GRAY
    r2 = p.add_run(event.get("pr_title", "No title"))
    r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = DARK

    # AI summary (green bar)
    if ai.get("pr_summary"):
        _add_left_bar_paragraph(doc, f"💡 {ai['pr_summary']}",
                                bar_color="10B981", bg_color="ECFDF5")

    # Info table
    rows = [
        ("PR Author",         event.get("pr_author", "N/A")),
        ("Merged By",         event.get("merged_by", "N/A")),
        ("Merged At",         event.get("merged_at", "N/A")),
        ("Merge Commit",      event.get("merge_commit_sha", "N/A")[:12] + "..."),
        ("Branch",            f"{event.get('source_branch', '?')} → {event.get('target_branch', '?')}"),
        ("Labels",            ", ".join(event.get("labels", [])) or "None"),
        ("Code Delta",        f"+{event.get('additions', 0)} / -{event.get('deletions', 0)} lines"),
        ("Commits",           str(event.get("commit_count", 0))),
        ("Files Changed",     str(len(event.get("files_changed", [])))),
        ("Approvals",         ", ".join(event.get("approvals", [])) or "None"),
        ("Changes Requested", ", ".join(event.get("changes_requested_by", [])) or "None"),
        ("Comments",          f"{event.get('comments', 0)} general  |  {event.get('review_comments', 0)} review"),
        ("Linked Issues",     str(len(event.get("linked_issues", [])))),
        ("PR URL",            event.get("pr_url", "N/A")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    for row_idx, (label, value) in enumerate(rows):
        lc = table.cell(row_idx, 0)
        vc = table.cell(row_idx, 1)
        _set_cell_bg(lc, "F9FAFB")
        _set_cell_border(lc)
        _set_cell_border(vc)
        lc.width = Inches(1.8)
        vc.width = Inches(5.7)

        lc.paragraphs[0].clear()
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True; lr.font.size = Pt(10); lr.font.color.rgb = DARK

        vc.paragraphs[0].clear()
        vr = vc.paragraphs[0].add_run(str(value))
        vr.font.size = Pt(10)

    # Commit messages
    commit_msgs = event.get("commit_messages", [])
    if commit_msgs:
        _add_paragraph(doc, "Commit Messages:", size=10, bold=True,
                       color=DARK, space_before=8, space_after=2)
        for msg in commit_msgs[:8]:
            _add_paragraph(doc, f"  • {msg}", size=9, color=GRAY, space_after=2)

    # AI: before/after
    if ai.get("before_after"):
        _add_paragraph(doc, "Before → After:", size=10, bold=True,
                       color=DARK, space_before=8, space_after=2)
        _add_left_bar_paragraph(doc, ai["before_after"], bar_color="3B82F6")

    # AI: key impacts
    if ai.get("key_impacts"):
        _add_paragraph(doc, "Key Impacts:", size=10, bold=True,
                       color=DARK, space_before=6, space_after=2)
        for impact in ai["key_impacts"]:
            _add_paragraph(doc, f"  • {impact}", size=9, color=GRAY, space_after=2)

    # AI: review notes
    if ai.get("review_notes"):
        _add_paragraph(doc, "Review Notes:", size=10, bold=True,
                       color=DARK, space_before=6, space_after=2)
        _add_paragraph(doc, ai["review_notes"], size=9, color=GRAY, space_after=4)

    # PR description
    if event.get("pr_body") and event["pr_body"] != "No description":
        _add_paragraph(doc, "PR Description:", size=10, bold=True,
                       color=DARK, space_before=6, space_after=2)
        _add_paragraph(doc, event["pr_body"][:400], size=9, color=GRAY, space_after=6)

    # Linked issues
    linked = event.get("linked_issues", [])
    if linked:
        _add_paragraph(doc, "Linked Issues:", size=10, bold=True,
                       color=DARK, space_before=6, space_after=4)
        for issue in linked:
            status = "Closed ✅" if issue["state"] == "closed" else "Open 🔴"
            _add_paragraph(doc,
                f"Issue #{issue['number']}: {issue['title']}  |  {status}",
                size=9, bold=True, color=PRIMARY, space_after=2)
            if issue.get("body") and issue["body"] != "No description":
                _add_paragraph(doc, f"  {issue['body'][:200]}",
                               size=9, italic=True, color=GRAY, space_after=4)

    # Files changed
    files = event.get("files_changed", [])
    if files:
        _add_paragraph(doc, "Files Changed:", size=10, bold=True,
                       color=DARK, space_before=6, space_after=2)
        for fname in files[:15]:
            _add_paragraph(doc, f"  {fname}", size=8, color=GRAY, space_after=1)

    # Code diffs
    diffs = event.get("file_diffs", [])
    if diffs:
        _add_paragraph(doc, "Code Diffs (sample):", size=10, bold=True,
                       color=DARK, space_before=6, space_after=2)
        for diff in diffs[:5]:
            if not diff.get("patch"):
                continue
            status_label = {
                "added": "ADDED", "modified": "MODIFIED",
                "removed": "REMOVED", "renamed": "RENAMED",
            }.get(diff.get("status"), "CHANGED")
            _add_paragraph(doc,
                f"{status_label}: {diff['filename']}  (+{diff.get('additions',0)} / -{diff.get('deletions',0)})",
                size=9, bold=True, color=PRIMARY, space_before=4, space_after=2)
            _add_paragraph(doc, diff["patch"][:600], size=8, color=GRAY, space_after=4)


def _build_footer(doc, generated_at):
    _add_divider(doc)
    _add_paragraph(doc,
        f"Merge Events Report  •  Generated by PR-Agent v2  •  {generated_at}",
        size=9, italic=True, color=LIGHT_GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)


# ─────────────────────────────────────────────
# SHARED STYLING HELPERS
# ─────────────────────────────────────────────

def _add_paragraph(doc, text, size=11, bold=False, italic=False,
                   color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_left_bar_paragraph(doc, text, bar_color="3B82F6", bg_color="EFF6FF"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(10)
    p.paragraph_format.left_indent   = Inches(0.3)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'),  '12')
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), bar_color)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def _add_divider(doc):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '4')
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), 'E5E7EB')
    pBdr.append(btm)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC"):
    tc       = cell._tc
    tcPr     = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)
